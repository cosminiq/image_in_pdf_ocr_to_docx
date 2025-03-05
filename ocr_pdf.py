from pdf2image import convert_from_path
import pytesseract
from docx import Document

# Numele fișierului PDF și DOCX
pdf_file = 'doc2.pdf'
output_docx = 'doc2.docx'

# Inițializăm documentul Word
document = Document()

# Convertim paginile PDF-ului în imagini
pages = convert_from_path(pdf_file)

# Parcurgem fiecare pagină și extragem textul
for i, page in enumerate(pages):
    text = pytesseract.image_to_string(page, lang='ron')  # Limba română
    document.add_heading(f'Pagina {i + 1}', level=1)
    document.add_paragraph(text)

# Salvăm textul într-un fișier DOCX
document.save(output_docx)

print(f"Textul a fost salvat în '{output_docx}'.")
